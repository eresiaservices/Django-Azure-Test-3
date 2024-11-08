# %%
#! pip install openai faiss-cpu langchain tiktoken PyPDF2 python-docxents.json

# %%
#! pip install --upgrade openai

# %%
#! pip install --upgrade openai httpx httpcore


# %%
import os
os.environ['KMP_DUPLICATE_LIB_OK'] = 'TRUE'
from openai import OpenAI
import openai
import faiss
import numpy as np
from langchain.embeddings import OpenAIEmbeddings
from langchain.docstore.document import Document



from langchain.schema import Document
import spacy.cli
spacy.cli.download("fr_core_news_sm")
import spacy
from docx import Document as DocxDocument
import PyPDF2
import fr_core_news_sm
import cv2
import pytesseract
from PIL import Image
import tiktoken





# %%

MODEL="gpt-4o"

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
clientstring = os.getenv("OPENAI_API_KEY")
openai.api_key = os.getenv("OPENAI_API_KEY")


# %%
import hashlib
import spacy
import tiktoken
import fitz  # PyMuPDF pour le traitement des PDF
from docx import Document as DocxDocument
import easyocr  # Utilisation d'EasyOCR pour l'OCR
from PIL import Image
import io



    # Charger le modèle de langue française de spacy
    
nlp = spacy.load("fr_core_news_sm")

    # Initialiser le tokenizer pour GPT-4o
tokenizer = tiktoken.encoding_for_model("gpt-4o")

    # Initialiser le lecteur EasyOCR pour le français
reader = easyocr.Reader(['fr'])

def preprocess_text(text):
        """Nettoie le texte avant segmentation."""
        text = text.strip()
        text = text.replace("\n", " ")
        return text

def split_into_sentences(text): 
        """Divise le texte en phrases à l'aide de spacy."""
        doc = nlp(text)
        sentences = [sent.text for sent in doc.sents]
        return sentences

def create_chunks(text, max_tokens=800):
        """Crée des chunks basés sur les phrases et une limite maximale de tokens."""
        sentences = split_into_sentences(text)
        current_chunk = []
        current_length = 0
        chunks = []

        for sentence in sentences:
            sentence_tokens = tokenizer.encode(sentence)
            sentence_length = len(sentence_tokens)
            if current_length + sentence_length <= max_tokens:
                current_chunk.append(sentence)
                current_length += sentence_length
            else:
                chunks.append(" ".join(current_chunk))
                current_chunk = [sentence]
                current_length = sentence_length

        if current_chunk:
            chunks.append(" ".join(current_chunk))

        return chunks

def extract_text_from_image(image_stream):
        """Extrait le texte d'une image en utilisant EasyOCR."""
        image = Image.open(image_stream)
        # Convertir l'image PIL en tableau numpy compatible avec EasyOCR
        image_np = np.array(image)
        # Utiliser EasyOCR pour l'OCR
        result = reader.readtext(image_np, detail=0)
        return " ".join(result)

def load_and_split_documents(docs_path):
        documents = []
        processed_images = set()  # Ensemble pour garder une trace des images traitées (hash MD5)

        # Vérifier si docs_path est un fichier unique ou un dossier
        if os.path.isfile(docs_path):
            # Traiter un fichier unique
            documents.extend(process_file(docs_path, processed_images))
        elif os.path.isdir(docs_path):
            # Parcourir tous les fichiers dans le dossier
            for filename in os.listdir(docs_path):
                filepath = os.path.join(docs_path, filename)
                if os.path.isfile(filepath):
                    documents.extend(process_file(filepath, processed_images))
        else:
            print(f"Le chemin {docs_path} n'est ni un fichier ni un dossier valide.")
        
        return documents

def process_file(filepath, processed_images):
        """
        Traite un fichier unique et le segmente.
        Retourne une liste d'objets Document (ou segments) à partir du fichier.
        """
        documents = []

        if filepath.endswith('.txt'):
            # Traitement des fichiers texte
            with open(filepath, 'r', encoding='utf-8') as f:
                text = f.read()
            chunks = create_chunks(text)
            for chunk in chunks:
                documents.append(Document(page_content=chunk))

        elif filepath.endswith('.pdf'):
            # Traitement des fichiers PDF
            with fitz.open(filepath) as pdf:
                for page_num, page in enumerate(pdf):
                    text = page.get_text() or ""
                    text = preprocess_text(text)

                    # Extraction des images du PDF
                    image_texts = []
                    for img_index, img in enumerate(page.get_images(full=True)):
                        xref = img[0]
                        base_image = pdf.extract_image(xref)
                        image_bytes = base_image["image"]

                        # Éviter de traiter la même image plusieurs fois
                        img_hash = hashlib.md5(image_bytes).hexdigest()
                        if img_hash in processed_images:
                            continue
                        processed_images.add(img_hash)

                        # OCR sur l'image
                        image_stream = io.BytesIO(image_bytes)
                        ocr_text = extract_text_from_image(image_stream)
                        image_texts.append(ocr_text)

                    combined_text = text + "\n" + "\n".join(image_texts)
                    chunks = create_chunks(preprocess_text(combined_text))
                    for chunk in chunks:
                        documents.append(Document(
                            page_content=chunk,
                            metadata={'page': page_num + 1, 'source': os.path.basename(filepath)}
                        ))

        elif filepath.endswith('.docx'):
            # Traitement des fichiers DOCX avec une approche dynamique
            doc = DocxDocument(filepath)
            elements = []
            rels = doc.part.rels

            # Parcourir les éléments du corps du fichier DOCX
            for child in doc.element.body:
                if child.tag == qn('w:p'):
                    # Extraction du texte des paragraphes
                    para = doc.paragraphs[len(elements)]
                    elements.append(('text', para.text))
                elif child.tag == qn('w:drawing'):
                    # Extraction des images
                    blip = child.xpath('.//a:blip', namespaces=child.nsmap)
                    if blip:
                        embed = blip[0].get(qn('r:embed'))
                        image_part = rels[embed]
                        image_data = image_part.target_part.blob

                        # Éviter de traiter les doublons d'images
                        img_hash = hashlib.md5(image_data).hexdigest()
                        if img_hash in processed_images:
                            continue
                        processed_images.add(img_hash)

                        # OCR sur l'image
                        image_stream = io.BytesIO(image_data)
                        ocr_text = extract_text_from_image(image_stream)
                        elements.append(('image', ocr_text))

            # Créer des chunks dynamiquement avec texte et images dans l'ordre d'apparition
            chunks = []
            current_chunk = ""
            current_length = 0

            for elem_type, content in elements:
                content = preprocess_text(content)
                content_tokens = tokenizer.encode(content)
                content_length = len(content_tokens)

                if current_length + content_length <= 800:
                    current_chunk += content + " "
                    current_length += content_length
                else:
                    chunks.append(current_chunk.strip())
                    current_chunk = content + " "
                    current_length = content_length

            if current_chunk:
                chunks.append(current_chunk.strip())

            for chunk in chunks:
                documents.append(Document(page_content=chunk))

        return documents

import shutil
import json

def add_docs(docs_path, treated_folder='C:/Users/guill/OneDrive/Documents/HOPE/Test_django/gestion_ressources_V2/Chatbot/Doc traités'):
        """
        Ajoute des documents ou des dossiers dans 'documents traités', segmente les nouveaux documents,
        et met à jour l'index FAISS avec les documents non encore traités.
        
        :param docs_path: str, chemin vers le dossier ou le fichier contenant les documents à ajouter.
        :param treated_folder: str, chemin vers le dossier contenant les documents déjà traités.
        """
        # Créer le dossier "documents traités" s'il n'existe pas
        if not os.path.exists(treated_folder):
            os.makedirs(treated_folder)

        # Liste des nouveaux documents à traiter
        new_docs = []

        # Vérifier si docs_path est un fichier unique ou un dossier
        if os.path.isfile(docs_path):
            filename = os.path.basename(docs_path)
            treated_filepath = os.path.join(treated_folder, filename)

            # Vérifier si le fichier existe déjà dans "documents traités"
            if os.path.exists(treated_filepath):
                print(f"Doc {filename} déjà traité")
            else:
                # Copier le document dans "documents traités"
                shutil.copy(docs_path, treated_folder)
                new_docs.append(docs_path)
                print(f"Ajout de {filename} dans {treated_folder}")

        elif os.path.isdir(docs_path):
            # Parcourir tous les fichiers dans le dossier
            for filename in os.listdir(docs_path):
                filepath = os.path.join(docs_path, filename)
                treated_filepath = os.path.join(treated_folder, filename)

                # Vérifier si le fichier existe déjà dans "documents traités"
                if os.path.exists(treated_filepath):
                    print(f"Doc {filename} déjà traité")
                else:
                    # Copier le document dans "documents traités"
                    shutil.copy(filepath, treated_folder)
                    new_docs.append(filepath)
                    print(f"Ajout de {filename} dans {treated_folder}")
        else:
            print(f"Le chemin {docs_path} n'est ni un fichier ni un dossier valide.")
            return

        # Si aucun nouveau document n'a été ajouté, arrêter ici
        if not new_docs:
            print("Aucun nouveau document à traiter.")
            return

        # Segmentation des nouveaux documents ajoutés
        print("Segmentation des nouveaux documents...")
        update_segments_json(treated_folder)

        # Mise à jour de l'index FAISS avec les nouveaux segments
        print("Mise à jour de l'index FAISS avec les nouveaux segments...")
        update_faiss_index_and_map()

def update_segments_json(docs_path, json_file_path='C:/Users/guill/OneDrive/Documents/HOPE/Test_django/gestion_ressources_V2/Chatbot/Segments/segments.json'):
        """
        Met à jour ou crée un fichier JSON des segments en ajoutant de nouveaux segments,
        tout en évitant les doublons. La fonction gère à la fois un fichier unique ou un dossier.
        
        :param docs_path: str, chemin vers le dossier ou fichier contenant les documents à segmenter.
        :param json_file_path: str, chemin vers le fichier JSON à mettre à jour ou créer.
        """
        # Liste des documents à segmenter
        documents = []

        # Vérifier si docs_path est un fichier ou un dossier
        if os.path.isfile(docs_path):
            # Si c'est un fichier unique, segmenter directement
            documents.extend(load_and_split_documents(docs_path))
        elif os.path.isdir(docs_path):
            # Si c'est un dossier, parcourir tous les fichiers à l'intérieur
            for filename in os.listdir(docs_path):
                filepath = os.path.join(docs_path, filename)
                if os.path.isfile(filepath):
                    documents.extend(load_and_split_documents(filepath))
        else:
            print(f"Le chemin {docs_path} n'est ni un fichier valide ni un dossier.")
            return

        # Créer la liste des nouveaux segments
        new_segments = []
        for doc in documents:
            segment_data = {
                "segment": doc.page_content,
                "metadata": doc.metadata if hasattr(doc, 'metadata') else {}
            }
            new_segments.append(segment_data)

        # Charger l'ancien fichier JSON s'il existe, sinon créer une liste vide
        if os.path.exists(json_file_path):
            with open(json_file_path, 'r') as json_file:
                old_segments = json.load(json_file)
        else:
            print(f"Le fichier {json_file_path} n'existe pas encore. Il va être créé.")
            old_segments = []  # Commencer avec une liste vide

        # Vérifier et éviter les doublons en comparant les métadonnées
        for new_segment in new_segments:
            if not any(old_seg['metadata'] == new_segment['metadata'] for old_seg in old_segments):
                old_segments.append(new_segment)  # Ajouter uniquement si non existant

        # Sauvegarder les segments mis à jour dans le fichier JSON
        with open(json_file_path, 'w') as json_file:
            json.dump(old_segments, json_file, indent=4)
        
        print(f"Le fichier JSON a été mis à jour ou créé avec succès dans {json_file_path}")

def update_faiss_index_and_map(embedding_model=OpenAIEmbeddings(model='text-embedding-ada-002', openai_api_key=clientstring), 
                                faiss_index_path='C:/Users/guill/OneDrive/Documents/HOPE/Test_django/gestion_ressources_V2/Chatbot/Index/faiss_index.bin', 
                                doc_id_map_path='C:/Users/guill/OneDrive/Documents/HOPE/Test_django/gestion_ressources_V2/Chatbot/Mapping/doc_id_map.json', 
                                json_file_path='C:/Users/guill/OneDrive/Documents/HOPE/Test_django/gestion_ressources_V2/Chatbot/Segments/segments.json'):
        """
        Met à jour l'index FAISS et le mapping doc_id_map en ajoutant de nouveaux segments.
        
        :param embedding_model: modèle utilisé pour générer les embeddings (par défaut OpenAI text-embedding-ada-002).
        :param faiss_index_path: str, chemin vers le fichier binaire de l'index FAISS.
        :param doc_id_map_path: str, chemin vers le fichier JSON du mapping doc_id_map.
        :param json_file_path: str, chemin vers le fichier JSON contenant les segments.
        """
        # Charger les segments depuis le fichier JSON
        with open(json_file_path, 'r') as json_file:
            segments_data = json.load(json_file)

        # Charger ou initialiser le doc_id_map
        if os.path.exists(doc_id_map_path):
            with open(doc_id_map_path, 'r') as json_file:
                try:
                    doc_id_map = json.load(json_file)
                except json.JSONDecodeError:
                    print(f"Erreur lors du chargement du mapping depuis {doc_id_map_path}. Le fichier est corrompu ou vide.")
                    doc_id_map = {}
        else:
            print(f"Mapping non trouvé à {doc_id_map_path}. Un nouveau mapping sera créé.")
            doc_id_map = {}

        # Ajouter de nouveaux documents
        new_document_embeddings = []
        new_doc_id_map = {}

        # Indexation à partir de la taille actuelle du doc_id_map
        starting_index = len(doc_id_map)
        for i, segment_info in enumerate(segments_data[starting_index:]):  # Traiter uniquement les nouveaux segments
            if not any(existing_seg['metadata'] == segment_info['metadata'] for existing_seg in doc_id_map.values()):
                segment_text = segment_info['segment']
                embedding = embedding_model.embed_documents([segment_text])
                new_document_embeddings.append(embedding[0])
                new_doc_id_map[starting_index + i] = segment_info

        # Ajouter les nouveaux embeddings à l'index FAISS
        if new_document_embeddings:
            new_embeddings_array = np.array(new_document_embeddings).astype('float32')

            # Si l'index FAISS existe déjà, on le charge, sinon on en crée un nouveau
            if os.path.exists(faiss_index_path):
                index = faiss.read_index(faiss_index_path)
                index.add(new_embeddings_array)
            else:
                embedding_dim = len(new_document_embeddings[0])
                index = faiss.IndexFlatL2(embedding_dim)
                index.add(new_embeddings_array)

            # Sauvegarder l'index FAISS
            faiss.write_index(index, faiss_index_path)
            print(f"Index FAISS mis à jour et sauvegardé dans {faiss_index_path}")

            # Mettre à jour le doc_id_map
            doc_id_map.update(new_doc_id_map)

            # Sauvegarder le mapping mis à jour
            with open(doc_id_map_path, 'w') as json_file:
                json.dump(doc_id_map, json_file, indent=4)
            print(f"Mapping doc_id_map mis à jour dans {doc_id_map_path}")
        else:
            print("Aucun nouveau segment à ajouter (tous les segments sont déjà présents).")

def remove_docs(file_paths, treated_folder='C:/Users/guill/OneDrive/Documents/HOPE/Test_django/gestion_ressources_V2/Chatbot/Doc traités', 
                    faiss_index_path='C:/Users/guill/OneDrive/Documents/HOPE/Test_django/gestion_ressources_V2/Chatbot/Index/faiss_index.bin', 
                    doc_id_map_path='C:/Users/guill/OneDrive/Documents/HOPE/Test_django/gestion_ressources_V2/Chatbot/Mapping/doc_id_map.json', 
                    json_file_path='C:/Users/guill/OneDrive/Documents/HOPE/Test_django/gestion_ressources_V2/Chatbot/Segments/segments.json'):
        """
        Supprime des fichiers de 'Documents traités', efface les fichiers segments/index/map,
        et recrée les segments, l'index FAISS et le mapping à partir des documents restants dans 'Documents traités'.
        
        :param file_paths: list, liste des chemins vers les fichiers à supprimer.
        :param treated_folder: str, chemin vers le dossier 'Documents traités'.
        :param faiss_index_path: str, chemin vers le fichier binaire de l'index FAISS.
        :param doc_id_map_path: str, chemin vers le fichier JSON du mapping doc_id_map.
        :param json_file_path: str, chemin vers le fichier JSON des segments.
        """
        # Supprimer les fichiers dans 'Documents traités'
        delete_files_in_treated_folder(file_paths, treated_folder)

        # Supprimer les fichiers segments/index/map
        print("Suppression des fichiers segments, index, et mapping...")
        if os.path.exists(faiss_index_path):
            os.remove(faiss_index_path)
            print(f"Fichier index FAISS supprimé : {faiss_index_path}")
        
        if os.path.exists(doc_id_map_path):
            os.remove(doc_id_map_path)
            print(f"Fichier mapping doc_id_map supprimé : {doc_id_map_path}")
        
        if os.path.exists(json_file_path):
            os.remove(json_file_path)
            print(f"Fichier segments JSON supprimé : {json_file_path}")

        # Vérifier s'il reste des documents dans 'Documents traités'
        if not os.listdir(treated_folder):
            print("Aucun document restant dans 'Documents traités'.")
            return

        # Recréer les segments et l'index à partir des documents restants
        print("Recréation des segments et de l'index FAISS à partir des documents restants...")
        treated_folder='C:/Users/guill/OneDrive/Documents/HOPE/Test_django/gestion_ressources_V2/Chatbot/Doc traités'
        update_segments_json(treated_folder)  # Mise à jour des segments
        update_faiss_index_and_map()  # Mise à jour de l'index FAISS et du mapping

def delete_files_in_treated_folder(file_paths, treated_folder='C:/Users/guill/OneDrive/Documents/HOPE/Test_django/gestion_ressources_V2/Chatbot/Doc traités'):
        """
        Supprime des fichiers du dossier 'Documents traités'. Gère aussi bien un fichier unique qu'un dossier de fichiers.
        
        :param file_paths: str ou list, chemin(s) vers le fichier ou le dossier à supprimer.
        :param treated_folder: str, chemin vers le dossier 'Documents traités'.
        """
        if not os.path.exists(treated_folder):
            print(f"Le dossier {treated_folder} n'existe pas.")
            return

        # Si un seul fichier est donné en paramètre, convertir en liste
        if isinstance(file_paths, str):
            file_paths = [file_paths]

        for file_path in file_paths:
            if os.path.isdir(file_path):
                # Si c'est un dossier, supprimer chaque fichier à l'intérieur
                for filename in os.listdir(file_path):
                    filepath = os.path.join(file_path, filename)
                    delete_single_file(filepath, treated_folder)
            else:
                # Si c'est un fichier unique, supprimer directement
                delete_single_file(file_path, treated_folder)

def delete_single_file(file_path, treated_folder):
        """
        Supprime un fichier unique dans 'Documents traités'.
        """
        filename = os.path.basename(file_path)  # Extraire le nom du fichier
        treated_filepath = os.path.join(treated_folder, filename)  # Chemin complet dans le dossier 'Documents traités'

        if os.path.exists(treated_filepath):
            try:
                os.remove(treated_filepath)
                print(f"Fichier supprimé : {treated_filepath}")
            except Exception as e:
                print(f"Erreur lors de la suppression de {treated_filepath} : {e}")
        else:
            print(f"Fichier non trouvé dans 'Documents traités' : {filename}")



#%%    
def get_relevant_documents(query, embedding_model=OpenAIEmbeddings(model='text-embedding-ada-002', openai_api_key=clientstring), 
                           faiss_index_path='C:/Users/guill/OneDrive/Documents/HOPE/Test_django/gestion_ressources_V2/Chatbot/Index/faiss_index.bin', 
                           doc_id_map_path='C:/Users/guill/OneDrive/Documents/HOPE/Test_django/gestion_ressources_V2/Chatbot/Mapping/doc_id_map.json', 
                           k=10):
    """
    Effectue une recherche des documents pertinents en fonction d'une requête textuelle.
    
    :param query: str, la requête pour laquelle on cherche les documents.
    :param embedding_model: modèle utilisé pour générer l'embedding de la requête (par défaut OpenAI text-embedding-ada-002).
    :param faiss_index_path: str, chemin vers le fichier binaire de l'index FAISS (par défaut 'C:/Users/guill/OneDrive/Documents/HOPE/Test_django/gestion_ressources_V2/Chatbot/Index/faiss_index.bin').
    :param doc_id_map_path: str, chemin vers le fichier JSON de mapping doc_id_map (par défaut 'C:/Users/guill/OneDrive/Documents/HOPE/Test_django/gestion_ressources_V2/Chatbot/Mapping/doc_id_map.json').
    :param k: int, nombre de documents pertinents à retourner (par défaut 10).
    :return: list, une liste des documents pertinents avec leurs métadonnées.
    """
    # Charger l'index FAISS
    index = faiss.read_index(faiss_index_path)

    # Charger le mapping des documents
    with open(doc_id_map_path, 'r') as json_file:
        doc_id_map = json.load(json_file)

    # Générer l'embedding de la requête
    query_embedding = embedding_model.embed_documents([query])[0]  # Générer l'embedding pour la requête
    query_embedding = np.array(query_embedding).astype('float32').reshape(1, -1)

    # Recherche des k documents les plus proches
    distances, indices = index.search(query_embedding, k)

    # Vérification si des documents ont été trouvés
    if len(indices[0]) == 0:
        return []

    # Récupérer les documents correspondants
    relevant_docs = [doc_id_map[str(idx)] for idx in indices[0]]  # Convertir l'indice en string pour le mapping
    return relevant_docs


# %%
def generate_answer(query, embedding_model=OpenAIEmbeddings(model='text-embedding-ada-002', 
                                                            openai_api_key=clientstring),
                    faiss_index_path='C:/Users/guill/OneDrive/Documents/HOPE/Test_django/gestion_ressources_V2/Chatbot/Index/faiss_index.bin', 
                    doc_id_map_path='C:/Users/guill/OneDrive/Documents/HOPE/Test_django/gestion_ressources_V2/Chatbot/Mapping/doc_id_map.json', 
                    model=MODEL):
    """
    Génère une réponse en fonction des documents pertinents trouvés ou à partir des connaissances du modèle.
    
    :param query: str, la requête de l'utilisateur.
    :param embedding_model: modèle utilisé pour générer l'embedding de la requête (par défaut OpenAI text-embedding-ada-002).
    :param faiss_index_path: str, chemin vers le fichier binaire de l'index FAISS (par défaut 'C:/Users/guill/OneDrive/Documents/HOPE/Test_django/gestion_ressources_V2/Chatbot/Index/faiss_index.bin').
    :param doc_id_map_path: str, chemin vers le fichier JSON de mapping doc_id_map (par défaut 'C:/Users/guill/OneDrive/Documents/HOPE/Test_django/gestion_ressources_V2/Chatbot/Mapping/doc_id_map.json').
    :param model: str, nom du modèle à utiliser pour générer la réponse.
    """
    # Trouver les documents pertinents
    relevant_docs = get_relevant_documents(query, embedding_model=embedding_model, 
                                           faiss_index_path=faiss_index_path, 
                                           doc_id_map_path=doc_id_map_path, k=10)
    
    
    # Créer le contexte et imprimer ce qui est envoyé
    if relevant_docs:
        context = "\n\n".join([doc['segment'] for doc in relevant_docs])
        prompt = f"""
Vous êtes un assistant spécialisé dans les services sociaux en entreprise, droits sociaux et la gestion des ressources humaines.

**Instructions :**

- Utilisez les informations fournies dans le contexte pour répondre de manière concise et détaillée à la question de l'utilisateur.
- Si nécessaire, vous pouvez répondre à la question en utilisant vos connaissances.
- Si vous ne trouvez pas de réponse pertinente dans le contexte, indiquez-le à l'utilisateur en disant : "Cette requête semble sortir du cadre de ma spécialisation, voici néanmoins une réponse." Et répondez à la question en utilisant vos connaissances.

**Contexte :**
{context}

**Question :**
{query}

**Réponse :**
"""
        
    else:
        prompt = f"""
Vous êtes un assistant spécialisé dans les services sociaux en entreprise, droits sociaux et la gestion des ressources humaines.

**Instructions :**

- Aucun document pertinent n'a été trouvé pour répondre à la question.
- Vous êtes autorisé à utiliser vos connaissances pour répondre de manière précise et détaillée.
- Indiquez à l'utilisateur que l'information provient de vos connaissances et non des documents fournis.

**Question :**
{query}

**Réponse :**
"""
        
    
    # Générer la réponse avec le modèle
    response = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": "Vous êtes un assistant spécialisé dans les services sociaux en entreprise, droits sociaux et la gestion des ressources humaines."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=1000,
        temperature=0.2,
    )


    answer = response.choices[0].message.content


    return answer
